#!/usr/bin/env python3
"""Document redlining tool — compare docs, read/write track changes, add comments.

Usage:
    python redline_tool.py compare <file1> <file2> [--output redline.docx] [--track-changes]
    python redline_tool.py read-changes <file.docx>
    python redline_tool.py suggest <file.docx> <changes.json> [--output suggested.docx]
    python redline_tool.py comment <file.docx> <comments.json> [--output commented.docx]

Files can be local paths, Google Drive file IDs, or Google Drive URLs.

Requires: python-docx, lxml (auto-installed if missing)
For Drive access: GOOGLE_CLIENT_ID, GOOGLE_CLIENT_SECRET, GOOGLE_REFRESH_TOKEN
"""

import difflib
import json
import os
import re
import sys
import tempfile
from datetime import datetime, timezone
from urllib.parse import parse_qs, urlparse


def _ensure_deps():
    """Auto-install python-docx and lxml if missing."""
    try:
        import docx  # noqa: F401
        import lxml  # noqa: F401
    except ImportError:
        import subprocess
        subprocess.check_call([sys.executable, '-m', 'pip', 'install', '-q',
                               'python-docx', 'lxml'])


def _is_drive_ref(path: str) -> bool:
    """Check if a path looks like a Drive ID or URL rather than a local file."""
    if os.path.isfile(path):
        return False
    if 'drive.google.com' in path or 'docs.google.com' in path:
        return True
    if re.match(r'^[a-zA-Z0-9_-]{20,}$', path):
        return True
    return False


def _extract_id(url_or_id: str) -> str:
    """Extract file ID from a Google Drive URL or return as-is."""
    if re.match(r'^[a-zA-Z0-9_-]+$', url_or_id):
        return url_or_id
    m = re.search(r'/d/([a-zA-Z0-9_-]+)', url_or_id)
    if m:
        return m.group(1)
    parsed = urlparse(url_or_id)
    qs = parse_qs(parsed.query)
    if 'id' in qs:
        return qs['id'][0]
    raise ValueError(f"Cannot extract file ID from: {url_or_id}")


sys.path.insert(0, os.path.join(os.path.dirname(os.path.abspath(__file__)), '..', '_shared'))
from google_auth import get_credentials


def _download_from_drive(file_ref: str) -> str:
    """Download a .docx from Google Drive to a temp file. Returns the local path."""
    from googleapiclient.discovery import build

    file_id = _extract_id(file_ref)
    creds = get_credentials()
    service = build('drive', 'v3', credentials=creds)

    meta = service.files().get(fileId=file_id, fields='name,mimeType').execute()
    mime = meta['mimeType']
    name = meta.get('name', 'document')

    if mime == 'application/vnd.google-apps.document':
        content = service.files().export(
            fileId=file_id,
            mimeType='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        ).execute()
    else:
        content = service.files().get_media(fileId=file_id).execute()

    suffix = '.docx'
    tmp = tempfile.NamedTemporaryFile(suffix=suffix, prefix=f"{name}_", delete=False)
    tmp.write(content)
    tmp.close()
    print(f"Downloaded: {name} -> {tmp.name}", file=sys.stderr)
    return tmp.name


def _resolve_file(path: str) -> str:
    """Resolve a file argument to a local path, downloading from Drive if needed."""
    if _is_drive_ref(path):
        return _download_from_drive(path)
    if not os.path.isfile(path):
        print(f"ERROR: File not found: {path}", file=sys.stderr)
        sys.exit(1)
    return path


def _get_output_path(args: list, default_name: str) -> str:
    """Extract --output path from args, or use OUTPUTS_PATH env var."""
    if '--output' in args:
        idx = args.index('--output')
        if idx + 1 < len(args):
            return args[idx + 1]
    output_dir = os.environ.get('OUTPUTS_PATH', '/mnt/user-data/outputs')
    return os.path.join(output_dir, default_name)


# ---------------------------------------------------------------------------
# Word XML namespace constants
# ---------------------------------------------------------------------------

W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
W_TAG = lambda tag: f'{{{W_NS}}}{tag}'


def _now_iso():
    return datetime.now(timezone.utc).strftime('%Y-%m-%dT%H:%M:%SZ')


# ---------------------------------------------------------------------------
# read-changes: Parse track changes from an existing document
# ---------------------------------------------------------------------------

def read_changes(file_path: str):
    """Read and display track changes (revisions) from a Word document."""
    import docx
    from lxml import etree

    local_file = _resolve_file(file_path)
    doc = docx.Document(local_file)

    body = doc.element.body
    changes = []

    # Find all insertion and deletion elements in document order
    for elem in body.iter():
        tag = etree.QName(elem.tag).localname if '}' in elem.tag else elem.tag

        if tag == 'ins':
            author = elem.get(W_TAG('author'), 'Unknown')
            date = elem.get(W_TAG('date'), '')
            # Collect inserted text from child runs
            text_parts = []
            for t in elem.iter(W_TAG('t')):
                if t.text:
                    text_parts.append(t.text)
            text = ''.join(text_parts)
            if text.strip():
                changes.append({
                    'type': 'insertion',
                    'text': text,
                    'author': author,
                    'date': date,
                })

        elif tag == 'del':
            author = elem.get(W_TAG('author'), 'Unknown')
            date = elem.get(W_TAG('date'), '')
            # Deletions use <w:delText> instead of <w:t>
            text_parts = []
            for dt in elem.iter(W_TAG('delText')):
                if dt.text:
                    text_parts.append(dt.text)
            text = ''.join(text_parts)
            if text.strip():
                changes.append({
                    'type': 'deletion',
                    'text': text,
                    'author': author,
                    'date': date,
                })

        elif tag == 'rPrChange':
            # Formatting change — note but less critical
            author = elem.get(W_TAG('author'), 'Unknown')
            date = elem.get(W_TAG('date'), '')
            changes.append({
                'type': 'format_change',
                'text': '(formatting change)',
                'author': author,
                'date': date,
            })

    # Also read comments for full context
    comments = _read_comments(doc)

    if not changes and not comments:
        print(f"No track changes or comments found in {os.path.basename(local_file)}")
        # Fall back to showing the document text so the agent has context
        print("\nDocument text:")
        for p in doc.paragraphs:
            if (p.text or '').strip():
                print(p.text)
        return

    if changes:
        print(f"Found {len(changes)} tracked change(s):\n")
        for i, c in enumerate(changes, 1):
            marker = '+' if c['type'] == 'insertion' else '-' if c['type'] == 'deletion' else '~'
            date_str = c['date'][:10] if c['date'] else ''
            print(f"  [{marker}] {c['type'].upper()} by {c['author']} ({date_str})")
            # Show text with context
            text = c['text']
            if len(text) > 200:
                text = text[:200] + '...'
            print(f"      {text}")
            print()

    if comments:
        print(f"\nFound {len(comments)} comment(s):\n")
        for c in comments:
            print(f"  [{c['author']}] {c['text']}")
            print()

    # Output as JSON for agent processing
    print("\n--- JSON ---")
    print(json.dumps({
        'changes': changes,
        'comments': comments,
        'total_changes': len(changes),
        'total_comments': len(comments),
        'insertions': len([c for c in changes if c['type'] == 'insertion']),
        'deletions': len([c for c in changes if c['type'] == 'deletion']),
    }, indent=2))

    if local_file.startswith(tempfile.gettempdir()):
        os.unlink(local_file)


def _read_comments(doc):
    """Extract comments from a document."""
    comments = []
    for rel in doc.part.rels.values():
        if 'comments' in rel.reltype:
            from lxml import etree
            comments_xml = etree.fromstring(rel.target_part.blob)
            for comment_el in comments_xml.iter(W_TAG('comment')):
                author = comment_el.get(W_TAG('author'), 'Unknown')
                date = comment_el.get(W_TAG('date'), '')
                text_parts = []
                for t in comment_el.iter(W_TAG('t')):
                    if t.text:
                        text_parts.append(t.text)
                text = ''.join(text_parts)
                if text.strip():
                    comments.append({
                        'author': author,
                        'date': date,
                        'text': text,
                    })
            break
    return comments


# ---------------------------------------------------------------------------
# suggest: Apply changes as proper Word track changes (w:ins / w:del)
# ---------------------------------------------------------------------------

def suggest(file_path: str, changes_json_path: str, output_path: str):
    """Apply suggested changes to a document as Word tracked revisions.

    The changes JSON should be an array of objects:
    [
      {
        "find": "original text to replace",
        "replace": "new suggested text",
        "author": "Jeeves"
      },
      {
        "find": "text to delete",
        "replace": "",
        "author": "Jeeves"
      }
    ]

    Produces a .docx with proper <w:del>/<w:ins> markup that shows up
    in Word's Track Changes review pane and can be accepted/rejected.
    """
    import docx
    from lxml import etree

    local_file = _resolve_file(file_path)

    with open(changes_json_path, 'r', encoding='utf-8') as f:
        changes_data = json.load(f)

    doc = docx.Document(local_file)
    now = _now_iso()
    applied = 0
    rev_id = 100  # Starting revision ID

    for change in changes_data:
        find_text = change.get('find', '')
        replace_text = change.get('replace', '')
        author = change.get('author', 'Jeeves')

        if not find_text:
            continue

        for para in doc.paragraphs:
            if find_text not in (para.text or ''):
                continue

            para_xml = para._element

            # Find the run(s) containing the target text
            # We need to handle text that may span multiple runs
            full_text = ''
            runs_with_positions = []
            for run_el in para_xml.iter(W_TAG('r')):
                for t_el in run_el.iter(W_TAG('t')):
                    if t_el.text:
                        start_pos = len(full_text)
                        full_text += t_el.text
                        runs_with_positions.append((run_el, t_el, start_pos, len(full_text)))

            find_start = full_text.find(find_text)
            if find_start == -1:
                continue

            find_end = find_start + len(find_text)

            # Build new XML elements to replace the affected runs
            # Strategy: rebuild the paragraph's runs with del/ins markup
            new_elements = []
            processed_up_to = 0

            for run_el, t_el, r_start, r_end in runs_with_positions:
                # Get run properties (formatting) to preserve
                rpr = run_el.find(W_TAG('rPr'))
                rpr_copy = None
                if rpr is not None:
                    rpr_copy = etree.tostring(rpr)

                text = t_el.text or ''
                local_start = max(find_start - r_start, 0)
                local_end = min(find_end - r_start, len(text))

                # Text before the change in this run
                before = text[:local_start] if r_start < find_start and local_start > 0 else ''
                # Text that's being changed in this run
                middle = text[max(local_start, 0):max(local_end, 0)] if r_start < find_end and r_end > find_start else ''
                # Text after the change in this run
                after = text[local_end:] if r_end > find_end and local_end < len(text) else ''

                # No overlap with change region — keep run as-is
                if r_end <= find_start or r_start >= find_end:
                    new_elements.append(('keep', run_el))
                    continue

                # Before text — keep as normal run
                if before:
                    r = etree.Element(W_TAG('r'))
                    if rpr_copy:
                        r.append(etree.fromstring(rpr_copy))
                    t = etree.SubElement(r, W_TAG('t'))
                    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
                    t.text = before
                    new_elements.append(('new', r))

                # Middle text — wrap in <w:del>
                if middle:
                    del_el = etree.Element(W_TAG('del'))
                    del_el.set(W_TAG('id'), str(rev_id))
                    del_el.set(W_TAG('author'), author)
                    del_el.set(W_TAG('date'), now)
                    rev_id += 1

                    r = etree.SubElement(del_el, W_TAG('r'))
                    if rpr_copy:
                        r.append(etree.fromstring(rpr_copy))
                    dt = etree.SubElement(r, W_TAG('delText'))
                    dt.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
                    dt.text = middle
                    new_elements.append(('new', del_el))

                    # Insert replacement text right after the deletion (only once, on last affected run)
                    if replace_text and r_end >= find_end:
                        ins_el = etree.Element(W_TAG('ins'))
                        ins_el.set(W_TAG('id'), str(rev_id))
                        ins_el.set(W_TAG('author'), author)
                        ins_el.set(W_TAG('date'), now)
                        rev_id += 1

                        r = etree.SubElement(ins_el, W_TAG('r'))
                        if rpr_copy:
                            r.append(etree.fromstring(rpr_copy))
                        t = etree.SubElement(r, W_TAG('t'))
                        t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
                        t.text = replace_text
                        new_elements.append(('new', ins_el))

                # After text — keep as normal run
                if after:
                    r = etree.Element(W_TAG('r'))
                    if rpr_copy:
                        r.append(etree.fromstring(rpr_copy))
                    t = etree.SubElement(r, W_TAG('t'))
                    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
                    t.text = after
                    new_elements.append(('new', r))

            # Now replace the runs in the paragraph XML
            # Remove old runs
            for run_el, _, _, _ in runs_with_positions:
                try:
                    para_xml.remove(run_el)
                except ValueError:
                    pass

            # Insert new elements (preserve non-run children like bookmarks, comments)
            insert_point = None
            for child in list(para_xml):
                tag = etree.QName(child.tag).localname if '}' in child.tag else child.tag
                if tag == 'pPr':
                    insert_point = child
                    break

            idx = list(para_xml).index(insert_point) + 1 if insert_point is not None else 0
            for elem_type, elem in new_elements:
                if elem_type == 'keep':
                    para_xml.insert(idx, elem)
                else:
                    para_xml.insert(idx, elem)
                idx += 1

            applied += 1
            break  # Only apply once per change entry

    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)

    print(f"Document with tracked changes saved to: {output_path}")
    print(f"Applied {applied} of {len(changes_data)} suggested change(s)")
    print("Open in Word and go to Review > Track Changes to accept/reject each suggestion.")

    if local_file.startswith(tempfile.gettempdir()):
        os.unlink(local_file)


# ---------------------------------------------------------------------------
# compare: Paragraph-level diff (visual or tracked changes)
# ---------------------------------------------------------------------------

def _extract_accepted_text(element):
    """Extract the 'accepted all changes' text from a docx XML element.

    Walks the XML tree and:
      - Skips <w:del> content entirely (deleted text is gone)
      - Includes <w:ins> content as normal text (insertions are accepted)
      - Includes normal <w:t> text as-is
    This gives you the document as it would look after accepting all tracked changes.
    """
    from lxml import etree

    parts = []
    # Track which elements to skip (children of <w:del>)
    skip_ancestors = set()

    for event, elem in etree.iterwalk(element, events=('start', 'end')):
        tag = etree.QName(elem.tag).localname if '}' in elem.tag else elem.tag

        if event == 'start':
            if tag == 'del':
                skip_ancestors.add(id(elem))
            elif tag == 'delText':
                # Always skip deleted text
                continue
            elif tag == 't':
                # Only include <w:t> text if we're not inside a <w:del>
                in_del = False
                parent = elem.getparent()
                while parent is not None:
                    ptag = etree.QName(parent.tag).localname if '}' in parent.tag else parent.tag
                    if ptag == 'del':
                        in_del = True
                        break
                    parent = parent.getparent()
                if not in_del and elem.text:
                    parts.append(elem.text)

    return ''.join(parts)


def _extract_all_accepted_blocks(doc):
    """Extract accepted text from ALL content — paragraphs AND table cells.

    Accepts all tracked changes (skips deletions, includes insertions).
    Returns a list of strings, one per text block.
    """
    from lxml import etree

    blocks = []
    body = doc.element.body
    for elem in body:
        tag = etree.QName(elem.tag).localname if '}' in elem.tag else elem.tag
        if tag == 'p':
            blocks.append(_extract_accepted_text(elem))
        elif tag == 'tbl':
            for row in elem.iter(W_TAG('tr')):
                row_texts = []
                for cell in row.iter(W_TAG('tc')):
                    cell_text = _extract_accepted_text(cell)
                    row_texts.append(cell_text)
                blocks.append('\t'.join(row_texts))
    return blocks


def compare(file1_path: str, file2_path: str, output_path: str, track_changes: bool = False):
    """Compare two .docx files and produce a redlined document.

    Both documents are read with all tracked changes accepted — deleted text
    is dropped, insertions are treated as normal text.  This means you can
    compare two 'commented' versions and see the net difference between them.
    Table cell content is included (critical for term sheets).
    """
    import docx
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_UNDERLINE

    local1 = _resolve_file(file1_path)
    local2 = _resolve_file(file2_path)

    doc1 = docx.Document(local1)
    doc2 = docx.Document(local2)

    # Extract accepted (clean) text from both docs, including tables
    paras1 = _extract_all_accepted_blocks(doc1)
    paras2 = _extract_all_accepted_blocks(doc2)

    if track_changes:
        _compare_with_track_changes(doc1, paras1, paras2, output_path)
    else:
        _compare_visual(local1, local2, paras1, paras2, output_path)

    # Clean up temp files
    for f in [local1, local2]:
        if f.startswith(tempfile.gettempdir()):
            os.unlink(f)


def _match_paragraphs(paras1, paras2, threshold=0.4):
    """Match paragraphs by similarity regardless of order.

    Returns:
        matches: list of (i, j, ratio) tuples — matched pairs
        unmatched1: set of indices in paras1 with no match (deletions)
        unmatched2: set of indices in paras2 with no match (insertions)
    """
    # Skip empty paragraphs for matching purposes
    used1 = set()
    used2 = set()
    matches = []

    # First pass: exact matches (fast)
    text_to_idx1 = {}
    for i, t in enumerate(paras1):
        if t.strip():
            text_to_idx1.setdefault(t, []).append(i)
    for j, t in enumerate(paras2):
        if t.strip() and t in text_to_idx1 and text_to_idx1[t]:
            i = text_to_idx1[t].pop(0)
            matches.append((i, j, 1.0))
            used1.add(i)
            used2.add(j)

    # Second pass: fuzzy match remaining non-empty paragraphs
    remaining1 = [(i, paras1[i]) for i in range(len(paras1))
                  if i not in used1 and paras1[i].strip()]
    remaining2 = [(j, paras2[j]) for j in range(len(paras2))
                  if j not in used2 and paras2[j].strip()]

    # Build similarity scores and greedily pick best matches
    scored = []
    for i, t1 in remaining1:
        for j, t2 in remaining2:
            ratio = difflib.SequenceMatcher(None, t1.split(), t2.split()).ratio()
            if ratio >= threshold:
                scored.append((ratio, i, j))
    scored.sort(reverse=True)

    for ratio, i, j in scored:
        if i not in used1 and j not in used2:
            matches.append((i, j, ratio))
            used1.add(i)
            used2.add(j)

    # Empty paragraphs: match by position (they're structural, not content)
    empty1 = [i for i in range(len(paras1)) if i not in used1 and not paras1[i].strip()]
    empty2 = [j for j in range(len(paras2)) if j not in used2 and not paras2[j].strip()]
    for i, j in zip(empty1, empty2):
        matches.append((i, j, 1.0))
        used1.add(i)
        used2.add(j)

    unmatched1 = set(range(len(paras1))) - used1
    unmatched2 = set(range(len(paras2))) - used2
    return matches, unmatched1, unmatched2


def _get_body_elements(doc):
    """Get top-level body elements (paragraphs and tables) in document order."""
    from lxml import etree
    elements = []
    for child in doc.element.body:
        tag = etree.QName(child.tag).localname if '}' in child.tag else child.tag
        if tag in ('p', 'tbl'):
            elements.append((tag, child))
    return elements


def _get_table_cell_texts(tbl_el):
    """Extract accepted text from each cell in a table, preserving row/col structure.

    Returns list of rows, each row is a list of cell texts (accepted, no tracked changes).
    """
    rows = []
    for row_el in tbl_el.iter(W_TAG('tr')):
        cells = []
        for cell_el in row_el.iter(W_TAG('tc')):
            cells.append(_extract_accepted_text(cell_el))
        rows.append(cells)
    return rows


def _diff_table_cell(cell_el, old_text, new_text, rev_id, author, date):
    """Apply word-level tracked changes to a single table cell.

    Replaces the cell's paragraph content with del/ins markup.
    Returns the updated rev_id.
    """
    from lxml import etree

    if old_text == new_text:
        return rev_id

    # Find the first paragraph in the cell to apply the diff to
    para_el = cell_el.find(W_TAG('p'))
    if para_el is None:
        return rev_id

    return _inline_diff_paragraph(para_el, old_text, new_text, rev_id, author, date)


def _compare_with_track_changes(base_doc, paras1, paras2, output_path):
    """Produce a document with proper Word track changes markup.

    Walks both documents' structure (paragraphs AND tables) and diffs them.
    For tables, diffs cell-by-cell to preserve table formatting.
    For paragraphs, uses similarity-based matching with word-level diffs.
    All text is read with tracked changes accepted first.
    """
    import docx
    from lxml import etree

    now = _now_iso()
    author = 'Jeeves'
    rev_id = 100
    changes = {'equal': 0, 'modified': 0, 'delete': 0, 'insert': 0}

    # Load doc2 to get its structure
    # (base_doc is doc1, we need doc2 for table structure comparison)
    # We get doc2 from the output_path context — but actually we need
    # both docs' XML. The caller passes paras1/paras2 as flat text lists.
    # We need to work with the XML directly.

    body = base_doc.element.body
    elems1 = _get_body_elements(base_doc)

    # For paragraph-only content, use the flat text matching approach
    # For tables, we diff cell-by-cell in the base doc's table structure

    # Separate paragraphs and tables
    para_elements = []
    para_texts = []
    table_idx = 0

    for tag, elem in elems1:
        if tag == 'p':
            para_elements.append(elem)
            para_texts.append(_extract_accepted_text(elem))
        elif tag == 'tbl':
            # Tables stay in place — we'll diff their cells directly
            pass

    # Match the flat paragraph texts against paras2's paragraph-only entries
    # But paras2 is a mixed list (paragraphs + table rows).
    # We need to identify which entries in paras1/paras2 are table rows vs paragraphs.
    # Since _extract_all_accepted_blocks uses \t to join table row cells,
    # we can distinguish: entries with \t are table rows.

    para_only_1 = [t for t in paras1 if '\t' not in t]
    para_only_2 = [t for t in paras2 if '\t' not in t]
    table_rows_1 = [t for t in paras1 if '\t' in t]
    table_rows_2 = [t for t in paras2 if '\t' in t]

    # --- Diff paragraphs (non-table content) ---
    matches, unmatched1, unmatched2 = _match_paragraphs(para_only_1, para_only_2)
    j_to_match = {}
    for i, j, ratio in matches:
        j_to_match[j] = (i, ratio)

    # Apply paragraph diffs to the base doc's paragraph elements
    para_idx = 0
    for tag, elem in elems1:
        if tag != 'p':
            continue
        old_text = _extract_accepted_text(elem)
        # Find this paragraph's index in para_only_1
        if para_idx < len(para_only_1):
            # Check if it's matched to a different para_only_2 entry
            matched_j = None
            for j, (i, ratio) in j_to_match.items():
                if i == para_idx:
                    matched_j = (j, ratio)
                    break
            if matched_j:
                j, ratio = matched_j
                if ratio < 0.999 and j < len(para_only_2):
                    rev_id = _inline_diff_paragraph(elem, old_text, para_only_2[j], rev_id, author, now)
                    changes['modified'] += 1
                else:
                    changes['equal'] += 1
            elif para_idx in unmatched1 and old_text.strip():
                _wrap_para_runs_in_del(elem, rev_id, author, now)
                rev_id += 1
                changes['delete'] += 1
            else:
                changes['equal'] += 1
        para_idx += 1

    # --- Diff tables cell-by-cell ---
    table_elements = [elem for tag, elem in elems1 if tag == 'tbl']

    # Match table rows between doc1 and doc2 by content
    for tbl_idx, tbl_el in enumerate(table_elements):
        rows1 = _get_table_cell_texts(tbl_el)
        # Get corresponding table rows from doc2's flat text
        # Each row in table_rows_2 is tab-separated cell text
        # We need to figure out which table_rows_2 entries correspond to this table

        row_elements = list(tbl_el.iter(W_TAG('tr')))

        for row_idx, row_el in enumerate(row_elements):
            if row_idx >= len(rows1):
                break
            old_cells = rows1[row_idx]

            # Find the matching row in table_rows_2
            old_row_text = '\t'.join(old_cells)

            # Search for the best matching row in table_rows_2
            best_match = None
            best_ratio = 0
            for tr2_idx, tr2_text in enumerate(table_rows_2):
                new_cells = tr2_text.split('\t')
                ratio = difflib.SequenceMatcher(None, old_row_text.split(), tr2_text.split()).ratio()
                if ratio > best_ratio:
                    best_ratio = ratio
                    best_match = (tr2_idx, new_cells)

            if best_match and best_ratio > 0.3:
                tr2_idx, new_cells = best_match
                # Remove from available pool so we don't re-match
                if tr2_idx < len(table_rows_2):
                    table_rows_2[tr2_idx] = ''  # Mark as used

                cell_elements = list(row_el.iter(W_TAG('tc')))
                for c_idx in range(min(len(cell_elements), len(old_cells), len(new_cells))):
                    old_c = old_cells[c_idx]
                    new_c = new_cells[c_idx]
                    if old_c != new_c:
                        rev_id = _diff_table_cell(cell_elements[c_idx], old_c, new_c, rev_id, author, now)
                        changes['modified'] += 1
                    else:
                        changes['equal'] += 1
            else:
                changes['equal'] += len(old_cells)

    # Add new paragraphs from doc2 that don't exist in doc1
    for j in sorted(unmatched2):
        if j < len(para_only_2) and para_only_2[j].strip():
            p_el = _make_ins_paragraph(para_only_2[j], rev_id, author, now)
            rev_id += 1
            body.append(p_el)
            changes['insert'] += 1

    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    base_doc.save(output_path)

    print(f"Redline (with track changes) saved to: {output_path}")
    print(f"Changes: {changes['equal']} unchanged, {changes['modified']} modified (inline diff), "
          f"{changes['delete']} deleted, {changes['insert']} inserted")
    print("Open in Word and enable Review > Track Changes to see all revisions.")


def _wrap_para_runs_in_del(para_el, rev_id, author, date):
    """Wrap all runs in a paragraph inside a <w:del> element."""
    from lxml import etree

    runs = list(para_el.iter(W_TAG('r')))
    if not runs:
        return

    del_el = etree.Element(W_TAG('del'))
    del_el.set(W_TAG('id'), str(rev_id))
    del_el.set(W_TAG('author'), author)
    del_el.set(W_TAG('date'), date)

    for run in runs:
        # Convert <w:t> to <w:delText>
        for t in run.iter(W_TAG('t')):
            new_dt = etree.Element(W_TAG('delText'))
            new_dt.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
            new_dt.text = t.text
            t.getparent().replace(t, new_dt)
        run.getparent().remove(run)
        del_el.append(run)

    # Insert del_el after pPr (or at start)
    ppr = para_el.find(W_TAG('pPr'))
    if ppr is not None:
        ppr.addnext(del_el)
    else:
        para_el.insert(0, del_el)


def _inline_diff_paragraph(para_el, old_text, new_text, rev_id, author, date):
    """Replace a paragraph's content with word-level tracked changes.

    Instead of deleting the entire paragraph and inserting a new one
    (which makes everything blue in Word), this does a word-level diff
    so only the changed words are marked as deletions/insertions.
    Unchanged words stay in their original formatting (black).
    """
    from lxml import etree

    old_words = old_text.split()
    new_words = new_text.split()
    matcher = difflib.SequenceMatcher(None, old_words, new_words)

    # Remove existing runs from the paragraph (keep pPr)
    ppr = para_el.find(W_TAG('pPr'))
    for child in list(para_el):
        if child.tag != W_TAG('pPr'):
            para_el.remove(child)

    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag == 'equal':
            # Unchanged words — plain run
            r = etree.SubElement(para_el, W_TAG('r'))
            t = etree.SubElement(r, W_TAG('t'))
            t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
            t.text = ' '.join(old_words[i1:i2]) + ' '

        elif tag == 'delete':
            # Deleted words — wrap in <w:del>
            del_el = etree.SubElement(para_el, W_TAG('del'))
            del_el.set(W_TAG('id'), str(rev_id))
            del_el.set(W_TAG('author'), author)
            del_el.set(W_TAG('date'), date)
            rev_id += 1
            r = etree.SubElement(del_el, W_TAG('r'))
            dt = etree.SubElement(r, W_TAG('delText'))
            dt.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
            dt.text = ' '.join(old_words[i1:i2]) + ' '

        elif tag == 'insert':
            # Inserted words — wrap in <w:ins>
            ins_el = etree.SubElement(para_el, W_TAG('ins'))
            ins_el.set(W_TAG('id'), str(rev_id))
            ins_el.set(W_TAG('author'), author)
            ins_el.set(W_TAG('date'), date)
            rev_id += 1
            r = etree.SubElement(ins_el, W_TAG('r'))
            t = etree.SubElement(r, W_TAG('t'))
            t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
            t.text = ' '.join(new_words[j1:j2]) + ' '

        elif tag == 'replace':
            # Changed words — delete old, insert new
            del_el = etree.SubElement(para_el, W_TAG('del'))
            del_el.set(W_TAG('id'), str(rev_id))
            del_el.set(W_TAG('author'), author)
            del_el.set(W_TAG('date'), date)
            rev_id += 1
            r = etree.SubElement(del_el, W_TAG('r'))
            dt = etree.SubElement(r, W_TAG('delText'))
            dt.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
            dt.text = ' '.join(old_words[i1:i2]) + ' '

            ins_el = etree.SubElement(para_el, W_TAG('ins'))
            ins_el.set(W_TAG('id'), str(rev_id))
            ins_el.set(W_TAG('author'), author)
            ins_el.set(W_TAG('date'), date)
            rev_id += 1
            r = etree.SubElement(ins_el, W_TAG('r'))
            t = etree.SubElement(r, W_TAG('t'))
            t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
            t.text = ' '.join(new_words[j1:j2]) + ' '

    return rev_id


def _make_ins_paragraph(text, rev_id, author, date):
    """Create a new paragraph element with <w:ins> wrapped run."""
    from lxml import etree

    p = etree.Element(W_TAG('p'))
    ins_el = etree.SubElement(p, W_TAG('ins'))
    ins_el.set(W_TAG('id'), str(rev_id))
    ins_el.set(W_TAG('author'), author)
    ins_el.set(W_TAG('date'), date)

    r = etree.SubElement(ins_el, W_TAG('r'))
    t = etree.SubElement(r, W_TAG('t'))
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    t.text = text
    return p


def _compare_visual(local1, local2, paras1, paras2, output_path):
    """Original visual comparison — red strikethrough / blue underline."""
    import docx
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_UNDERLINE

    redline = docx.Document()

    header_para = redline.add_paragraph()
    run = header_para.add_run("REDLINE COMPARISON")
    run.bold = True
    run.font.size = Pt(14)
    header_para = redline.add_paragraph()
    run = header_para.add_run(f"Base: {os.path.basename(local1)}")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(128, 128, 128)
    header_para = redline.add_paragraph()
    run = header_para.add_run(f"Revised: {os.path.basename(local2)}")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(128, 128, 128)
    redline.add_paragraph()

    matcher = difflib.SequenceMatcher(None, paras1, paras2)
    changes = {'equal': 0, 'delete': 0, 'insert': 0, 'replace': 0}

    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag == 'equal':
            for text in paras1[i1:i2]:
                p = redline.add_paragraph()
                run = p.add_run(text)
                run.font.color.rgb = RGBColor(0, 0, 0)
            changes['equal'] += (i2 - i1)

        elif tag == 'delete':
            for text in paras1[i1:i2]:
                p = redline.add_paragraph()
                run = p.add_run(text)
                run.font.color.rgb = RGBColor(255, 0, 0)
                run.font.strikethrough = True
            changes['delete'] += (i2 - i1)

        elif tag == 'insert':
            for text in paras2[j1:j2]:
                p = redline.add_paragraph()
                run = p.add_run(text)
                run.font.color.rgb = RGBColor(0, 0, 255)
                run.font.underline = WD_UNDERLINE.SINGLE
            changes['insert'] += (j2 - j1)

        elif tag == 'replace':
            for text in paras1[i1:i2]:
                p = redline.add_paragraph()
                run = p.add_run(text)
                run.font.color.rgb = RGBColor(255, 0, 0)
                run.font.strikethrough = True
            for text in paras2[j1:j2]:
                p = redline.add_paragraph()
                run = p.add_run(text)
                run.font.color.rgb = RGBColor(0, 0, 255)
                run.font.underline = WD_UNDERLINE.SINGLE
            changes['replace'] += max(i2 - i1, j2 - j1)

    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    redline.save(output_path)

    print(f"Redline saved to: {output_path}")
    print(f"Changes: {changes['delete']} deletions, {changes['insert']} insertions, "
          f"{changes['replace']} replacements, {changes['equal']} unchanged paragraphs")


# ---------------------------------------------------------------------------
# comment: Add Word comments (unchanged from before)
# ---------------------------------------------------------------------------

def comment(file_path: str, comments_json_path: str, output_path: str):
    """Add Word comments to a .docx at matching paragraphs."""
    import docx
    from lxml import etree

    local_file = _resolve_file(file_path)

    with open(comments_json_path, 'r', encoding='utf-8') as f:
        comments_data = json.load(f)

    doc = docx.Document(local_file)

    nsmap = {'w': W_NS}

    from docx.opc.part import Part
    from docx.opc.packuri import PackURI

    comments_part = None
    comments_element = None

    for rel in doc.part.rels.values():
        if 'comments' in rel.reltype:
            comments_part = rel.target_part
            comments_element = etree.fromstring(comments_part.blob)
            break

    if comments_element is None:
        comments_element = etree.Element(f'{{{W_NS}}}comments', nsmap=nsmap)

    comment_id = 0
    matched = 0
    now = _now_iso()

    for entry in comments_data:
        para_match = entry.get('paragraph_match', '')
        comment_text = entry.get('comment', '')
        author = entry.get('author', 'Jeeves')

        for para in doc.paragraphs:
            if para_match.lower() in (para.text or '').lower():
                comment_el = etree.SubElement(comments_element, f'{{{W_NS}}}comment')
                comment_el.set(f'{{{W_NS}}}id', str(comment_id))
                comment_el.set(f'{{{W_NS}}}author', author)
                comment_el.set(f'{{{W_NS}}}date', now)

                cp = etree.SubElement(comment_el, f'{{{W_NS}}}p')
                cr = etree.SubElement(cp, f'{{{W_NS}}}r')
                ct = etree.SubElement(cr, f'{{{W_NS}}}t')
                ct.text = comment_text

                para_xml = para._element
                range_start = etree.Element(f'{{{W_NS}}}commentRangeStart')
                range_start.set(f'{{{W_NS}}}id', str(comment_id))
                para_xml.insert(0, range_start)

                range_end = etree.SubElement(para_xml, f'{{{W_NS}}}commentRangeEnd')
                range_end.set(f'{{{W_NS}}}id', str(comment_id))

                ref_run = etree.SubElement(para_xml, f'{{{W_NS}}}r')
                ref_rpr = etree.SubElement(ref_run, f'{{{W_NS}}}rPr')
                ref_style = etree.SubElement(ref_rpr, f'{{{W_NS}}}rStyle')
                ref_style.set(f'{{{W_NS}}}val', 'CommentReference')
                comment_ref = etree.SubElement(ref_run, f'{{{W_NS}}}commentReference')
                comment_ref.set(f'{{{W_NS}}}id', str(comment_id))

                comment_id += 1
                matched += 1
                break
        else:
            print(f"WARNING: No paragraph match for: '{para_match[:60]}...'", file=sys.stderr)

    if comments_part is not None:
        comments_part._blob = etree.tostring(comments_element, xml_declaration=True, encoding='UTF-8', standalone=True)
    else:
        comments_blob = etree.tostring(comments_element, xml_declaration=True, encoding='UTF-8', standalone=True)
        comments_part_uri = PackURI('/word/comments.xml')
        content_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml'
        comments_part = Part(
            comments_part_uri,
            content_type,
            comments_blob,
            doc.part.package,
        )
        doc.part.relate_to(comments_part, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments')

    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)

    print(f"Commented document saved to: {output_path}")
    print(f"Comments added: {matched} of {len(comments_data)}")

    if local_file.startswith(tempfile.gettempdir()):
        os.unlink(local_file)


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------

# ---------------------------------------------------------------------------
# drive-comment: Add comments directly to a Google Doc via Drive API
# ---------------------------------------------------------------------------

def drive_comment(file_ref: str, comments_json_path: str):
    """Add comments directly to a Google Doc using the Drive Comments API.

    This adds real Google Doc comments (visible in the sidebar) anchored to
    specific text — no download/re-upload needed.  The original doc is
    modified in place.

    The comments JSON should be an array of objects:
    [
      {
        "anchor_text": "text to anchor the comment to",
        "comment": "The comment content",
        "author": "Jeeves"
      }
    ]

    anchor_text is used to find the text in the doc and attach the comment
    to that region.  If anchor_text is empty or omitted, the comment is
    added as an unanchored (file-level) comment.
    """
    from googleapiclient.discovery import build

    file_id = _extract_id(file_ref)
    creds = get_credentials()
    service = build('drive', 'v3', credentials=creds)

    with open(comments_json_path, 'r', encoding='utf-8') as f:
        comments_data = json.load(f)

    added = 0
    for entry in comments_data:
        anchor_text = entry.get('anchor_text', '').strip()
        comment_text = entry.get('comment', '').strip()
        if not comment_text:
            continue

        body = {'content': comment_text}

        # quotedFileContent anchors the comment to specific text in the doc.
        # The value must be an EXACT substring of the document's plain text.
        # mimeType must be text/plain for Google Docs.
        if anchor_text:
            body['quotedFileContent'] = {
                'mimeType': 'text/plain',
                'value': anchor_text,
            }

        try:
            result = service.comments().create(
                fileId=file_id,
                body=body,
                fields='id,content,quotedFileContent,anchor',
            ).execute()
            added += 1
            anchored = "anchored" if result.get('anchor') else "file-level (text not found in doc)"
            label = f' on "{anchor_text[:60]}"' if anchor_text else ''
            print(f"  [{anchored}] Added comment{label}: {comment_text[:80]}")
        except Exception as e:
            print(f"  WARNING: Failed to add comment: {e}", file=sys.stderr)

    print(f"\nDone — added {added} of {len(comments_data)} comment(s) to the Google Doc.")
    print(f"View at: https://docs.google.com/document/d/{file_id}/edit")


def main():
    if hasattr(sys.stdout, 'reconfigure'):
        sys.stdout.reconfigure(encoding='utf-8', errors='replace')
    if hasattr(sys.stderr, 'reconfigure'):
        sys.stderr.reconfigure(encoding='utf-8', errors='replace')

    args = sys.argv[1:]
    if not args or args[0] in ('-h', '--help'):
        print("Usage:", file=sys.stderr)
        print("  python redline_tool.py compare <file1> <file2> [--track-changes] [--output redline.docx]", file=sys.stderr)
        print("  python redline_tool.py read-changes <file.docx>", file=sys.stderr)
        print("  python redline_tool.py suggest <file.docx> <changes.json> [--output suggested.docx]", file=sys.stderr)
        print("  python redline_tool.py comment <file.docx> <comments.json> [--output commented.docx]", file=sys.stderr)
        print("  python redline_tool.py drive-comment <google_doc_id_or_url> <comments.json>", file=sys.stderr)
        sys.exit(1)

    _ensure_deps()

    command = args[0]

    if command == 'compare':
        if len(args) < 3:
            print("ERROR: compare requires two file arguments", file=sys.stderr)
            sys.exit(1)
        output = _get_output_path(args, 'redline_output.docx')
        tc = '--track-changes' in args
        compare(args[1], args[2], output, track_changes=tc)

    elif command == 'read-changes':
        if len(args) < 2:
            print("ERROR: read-changes requires a .docx file", file=sys.stderr)
            sys.exit(1)
        read_changes(args[1])

    elif command == 'suggest':
        if len(args) < 3:
            print("ERROR: suggest requires a .docx file and a changes JSON file", file=sys.stderr)
            sys.exit(1)
        output = _get_output_path(args, 'suggested_output.docx')
        suggest(args[1], args[2], output)

    elif command == 'comment':
        if len(args) < 3:
            print("ERROR: comment requires a .docx file and a comments JSON file", file=sys.stderr)
            sys.exit(1)
        output = _get_output_path(args, 'commented_output.docx')
        comment(args[1], args[2], output)

    elif command == 'drive-comment':
        if len(args) < 3:
            print("ERROR: drive-comment requires a Google Doc ID/URL and a comments JSON file", file=sys.stderr)
            sys.exit(1)
        drive_comment(args[1], args[2])

    else:
        print(f"ERROR: Unknown command '{command}'. Use 'compare', 'read-changes', 'suggest', 'comment', or 'drive-comment'.", file=sys.stderr)
        sys.exit(1)


if __name__ == '__main__':
    main()
